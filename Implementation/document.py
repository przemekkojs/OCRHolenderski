import os

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from formatter import coordinates_to_lines
from complex_types import row, layout_row
from languages import language_code_to_adjective

def __remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr

    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'nil')
        tblBorders.append(border)

    tblPr.append(tblBorders)

def create_document(contents:list[row], output_path:str, lang_from:str, lang_to:str, debug:bool = False) -> None:
    if debug:
        print("Rozpoczęcie tworzenia dokumentu programu MS Word")

    lines:list[layout_row] = coordinates_to_lines(contents, debug)

    doc = Document()

    # === NAGŁÓWEK ===
    section = doc.sections[0]
    header = section.header
    header_contents = header.paragraphs[0]
    header_run = header_contents.add_run(f"Uwierzytelnione tłumaczenie z języka {language_code_to_adjective(lang_from)}")
    header_run.italic = True

    # === LINIA NAGŁÓWKA ===
    p_hc = header_contents._element
    p_hc_Pr = p_hc.get_or_add_pPr()
    p_hc_Border = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')

    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')

    p_hc_Border.append(bottom)
    p_hc_Pr.append(p_hc_Border)
    # ===
    # ===

    last_y:int = 0
    buffer:list[str] = []
    table = None
    table_cols = 0

    font_size:int = 11
    alignment:str = "L"

    for line in lines:
        translated:str = line.translation
        bad_translation:bool = False

        if translated == "":
            translated = line.word
            bad_translation = True

        cur_y:int = line.document_row

        if cur_y != last_y:
            if len(buffer) == 1:
                if table != None:                    
                    table = None
                    doc.add_paragraph()                   

                paragraph = doc.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(0)

                match alignment:
                    case 'L':
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    case 'C':
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    case 'R':
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    case _:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT               

                run = paragraph.add_run(buffer[0].strip())

                if bad_translation:
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                else:
                    run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

                run.font.size = Pt(11) # Pt(font_size)

            else:
                row_length = len(buffer)

                if table is None or row_length != table_cols:
                    if table is None:
                        paragraph = doc.add_paragraph()

                    paragraph.paragraph_format.space_after = Pt(0)

                    table = doc.add_table(rows=1, cols=row_length)
                    __remove_table_borders(table)
                    table.style = 'Table Grid'
                    table_cols = row_length
                    row_cells = table.rows[0].cells
                else:
                    row_cells = table.add_row().cells

                for index, text in enumerate(buffer):
                    row_cells[index].text = text.strip()                
            
            buffer.clear()

        buffer.append(translated)
        last_y = cur_y
        font_size:int = line.font_size
        alignment:str = line.alignment

    if buffer:
        if len(buffer) == 1:
            if table != None:                    
                table = None
                doc.add_paragraph()

            paragraph = doc.add_paragraph()

            match alignment:
                case 'L':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                case 'C':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                case 'R':
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                case _:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT               

            run = paragraph.add_run(buffer[0].strip())

            if bad_translation:
                    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            else:
                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)

            run.font.size = Pt(11) # Pt(font_size)
        else:
            row_length = len(buffer)

            if table is None or row_length != table_cols:
                if table is None:
                    table = doc.add_table(rows=1, cols=row_length)

                __remove_table_borders(table)
                table.style = 'Table Grid'
                table_cols = row_length
                row_cells = table.rows[0].cells
            else:
                row_cells = table.add_row().cells

            for index, text in enumerate(buffer):
                row_cells[index].text = text.strip()

    doc.save(output_path)

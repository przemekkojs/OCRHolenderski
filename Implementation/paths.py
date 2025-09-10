from datetime import datetime
from docx import Document
from typing import overload

import os
import json
import aiofiles

BASE_DIR:str = os.path.dirname(__file__)

async def __get_dictionary(lang_from:str, lang_to:str) -> dict[str, str]:
    dictionary_path:str = os.path.join(load_path('DICT_FOLDER'), f"{lang_from}_{lang_to}.json")
            
    if not os.path.exists(dictionary_path):
        async with aiofiles.open(dictionary_path, 'w', encoding='utf-8') as file:
            await file.write("{}")
        print(f'Utworzono plik słownika >> {dictionary_path} <<')
        return {}

    async with aiofiles.open(dictionary_path, 'r', encoding='utf-8') as f:
        content = await f.read()
        try:
            dictionary: dict[str, str] = json.loads(content)
        except json.JSONDecodeError:
            dictionary = {}

    print(f'Plik słownika >> {dictionary_path} << już istnieje')
    return dictionary

async def get_phrase_from_dictionary(phrase:str, lang_from:str, lang_to:str) -> str:
    dictionary = await __get_dictionary(lang_from, lang_to)
    return dictionary.get(phrase, "")

async def set_phrase_to_dictionary(phrase:str, translated:str, lang_from:str, lang_to:str) -> str:
    print(f"Zapisywanie >> {phrase} << do słownika ({lang_from} => {lang_to})")

    dictionary = await __get_dictionary(lang_from, lang_to)
    dictionary[phrase] = translated
    path = os.path.join(load_path('DICT_FOLDER'), f"{lang_from}_{lang_to}.json")    

    async with aiofiles.open(path, "w", encoding="utf-8") as f:
        await f.write(json.dumps(dictionary, ensure_ascii=False, indent=4))

    return translated

def load_path(key:str) -> None:
    with open("paths.json", "r", encoding="utf-8") as f:
        config = json.load(f)

    return os.path.normpath(config[key])

def __save(path:str, file_name:str, contents, debug:bool=False) -> None:
    full_path:str = os.path.join(path, file_name)
    os.makedirs(os.path.dirname(full_path), exist_ok=True)

    _, extension = os.path.splitext(file_name)

    try:
        if extension == ".docx":
            doc = Document(full_path) if os.path.exists(full_path) else Document()
            doc.add_paragraph(contents)
            doc.save(full_path)
        else:        
            with open(full_path, 'w', encoding='utf-8') as file:
                file.write(contents)

        if debug:
            print(f"Zapisano >> {file_name} << do lokalizacji >> {full_path} <<")
    except Exception as e:
        save_logs(str(e))

        if debug:
            print(f"Zapisywanie >> {file_name} << do lokalizacji >> {full_path} << ZAKOŃCZONE NIEPOWODZENIEM")
            print(e)


@overload
def save_any(full_path: str, contents) -> None: ...
@overload
def save_any(path: str, file_name: str, contents) -> None: ...

def save_any(*args):
    if len(args) == 2:
        full_path, contents = args
        path:str = os.path.dirname(full_path)
        file_name:str = os.path.basename(full_path)
        __save(path, file_name, contents)
    elif len(args) == 3:
        path, file_name, contents = args        
        __save(path, file_name, contents)
    else:
        raise TypeError("save_any() expects 2 or 3 arguments")

def save_logs(contents) -> None:
    now = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    __save(load_path("LOGS_FOLDER"), f"{now}.log", contents)

def save_tmp(file_name:str, contents) -> None:
    __save(load_path("TMP_FOLDER"), file_name, contents)

def save_prefs(file_name:str, contents) -> None:
    __save(load_path("PREFS_FOLDER"), file_name, contents)

def save_model(key:str, model, tokenizer) -> None:
    models_path:str = load_path("MODELS_FOLDER")
    cur_dir:str = os.path.join(models_path, key)

    model.save_pretrained(cur_dir)
    tokenizer.save_pretrained(cur_dir)

def model_exists(key:str) -> bool:
    models_path:str = load_path("MODELS_FOLDER")
    cur_dir:str = os.path.join(models_path, key)

    return os.path.exists(cur_dir)

def model_path(key:str) -> str:
    models_path:str = load_path("MODELS_FOLDER")
    return os.path.join(models_path, key)

def filter_path(name:str) -> str:
    path:str = os.path.join('./Filters', name)

    if not os.path.exists(path):
        os.mkdir(path)

    return path
